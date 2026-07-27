#!/usr/bin/env python3
"""Three Chapter 3 methods tables (sampling and validation design) from the GLKN exports, rendered in
a clean thesis style (top rule, header rule, bottom rule, no vertical rules). This is a formatting
pass: the only computed values are the specified unit conversions and the detection-rate ratio.

Tables (each written as a tidy csv, a viewing png, and an editable docx):
  A  per-agent change detection rate by grid cell size (detection rate = n_with_change/n_cells_complete)
  B  counts of complete cells containing each agent, by grid cell size (the absolute-count companion)
  C  disturbance polygon size by agent (areas converted to hectares, total to square kilometers)

Inputs (globbed from reports/GLKN_change_agents/, from the current GLKN Drive folders only):
  glkn_grid_proportions_per_agent_5070_4agent.csv   grid cell-size analysis (four agents)
  glkn_eda_changeagents_<year>.csv                  per-year change-agent polygon-area summaries

Run: python scripts/build_chapter3_tables.py
Requires: pandas, matplotlib, python-docx
"""

import glob
import os
import textwrap

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INDIR = os.path.join(REPO, "reports", "GLKN_change_agents")
OUTDIR = os.path.join(REPO, "manuscript_formatting", "tables")

AGENTS = ["harvest", "development", "beaver", "insect_disease_mort"]     # four change classes
DISPLAY = {"harvest": "Harvest", "development": "Development",
           "beaver": "Beaver", "insect_disease_mort": "Insect/Disease"}
SEL_PX = 112                                                            # selected cell size


def find(pattern):
    hits = [p for p in glob.glob(os.path.join(INDIR, pattern)) if "/.git/" not in p]
    if not hits:
        raise SystemExit(f"input not found: {pattern} in {INDIR}")
    return sorted(hits)[0]


# ---------------------------------------------------------------- clean png renderer
def render_png(path, title, subtitle, caption, headers, rows, aligns, bold_rows=(), footnote=None):
    # auto column widths from the longest header line or cell string, so nothing overlaps
    ncol, nrow = len(headers), len(rows)
    char_w, pad = 0.083, 0.30
    hdr_lines = [h.split("\n") for h in headers]
    col_w = [max(len(s) for s in (hdr_lines[j] + [r[j] for r in rows])) * char_w + pad
             for j in range(ncol)]
    left = 0.4
    right = left + sum(col_w)
    fig_w = right + 0.4
    row_in, line_h = 0.34, 0.27
    max_hlines = max(len(h) for h in hdr_lines)
    header_in = max_hlines * line_h + 0.06
    capw = max(40, int((fig_w - 0.8) / char_w))
    cap_lines = len(textwrap.wrap(caption, capw)) if caption else 0
    foot_lines = len(textwrap.wrap(footnote, capw)) if footnote else 0
    title_in = 0.36
    sub_in = 0.26 if subtitle else 0.0
    cap_in = cap_lines * 0.20 + (0.10 if caption else 0)
    foot_in = foot_lines * 0.18 + (0.14 if footnote else 0)
    fig_h = 0.2 + title_in + sub_in + cap_in + 0.12 + header_in + nrow * row_in + foot_in + 0.15

    fig = plt.figure(figsize=(fig_w, fig_h))
    ax = fig.add_axes([0, 0, 1, 1]); ax.set_xlim(0, fig_w); ax.set_ylim(0, fig_h)
    ax.invert_yaxis(); ax.axis("off")
    y = 0.22
    ax.text(left, y, title, fontsize=14, fontweight="bold", va="top"); y += title_in
    if subtitle:
        ax.text(left, y, subtitle, fontsize=11, va="top", color="0.25"); y += sub_in
    if caption:
        ax.text(left, y, "\n".join(textwrap.wrap(caption, capw)), fontsize=9, va="top", color="0.3")
        y += cap_in
    y += 0.12

    edges = [left]
    for w in col_w:
        edges.append(edges[-1] + w)

    def cellx(j):
        return (edges[j] + 0.10, "left") if aligns[j] == "l" else (edges[j + 1] - 0.10, "right")

    top_rule = y
    ax.plot([left, right], [top_rule, top_rule], color="0.1", lw=1.4)
    for j, lines in enumerate(hdr_lines):                              # multi-line headers, bottom-aligned
        x, ha = cellx(j)
        for k, ln in enumerate(reversed(lines)):
            ax.text(x, top_rule + header_in - (k + 0.5) * line_h, ln, fontsize=10.3, va="center",
                    ha=ha, fontweight="bold")
    y = top_rule + header_in
    ax.plot([left, right], [y, y], color="0.1", lw=0.9)                # header rule
    for i, r in enumerate(rows):
        for j, txt in enumerate(r):
            x, ha = cellx(j)
            ax.text(x, y + row_in / 2, txt, fontsize=10.3, va="center", ha=ha,
                    fontweight="bold" if i in bold_rows else "normal")
        y += row_in
    ax.plot([left, right], [y, y], color="0.1", lw=1.4)                # bottom rule
    if footnote:
        y += 0.16
        ax.text(left, y, "\n".join(textwrap.wrap(footnote, capw)), fontsize=8.5, va="top", color="0.35")

    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    print(f"wrote {path}")


# ---------------------------------------------------------------- clean docx renderer
def _set_border(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    for edge, on in edges.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}"); borders.append(el)
        if on:
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(on))
            el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


def _para(doc, text, size, bold=False, italic=False, color=None, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def build_docx(path, title, subtitle, caption, headers, rows, aligns, bold_rows=(), footnote=None):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    _para(doc, title, 13, bold=True, after=2)
    if subtitle:
        _para(doc, subtitle, 11, italic=True, color=(0x40, 0x40, 0x40), after=4)
    if caption:
        _para(doc, caption, 10, color=(0x33, 0x33, 0x33), after=6)

    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.cell(0, j); c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT if aligns[j] == "l" else WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(h.replace("\n", " ")); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
        _set_border(c, {"top": 12, "bottom": 6, "left": 0, "right": 0})
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i + 1, j)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT if aligns[j] == "l" else WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(val); r.font.name = "Times New Roman"; r.font.size = Pt(11)
            r.bold = i in bold_rows
            edges = {"top": 0, "bottom": 12 if i == len(rows) - 1 else 0, "left": 0, "right": 0}
            _set_border(c, edges)
    if footnote:
        _para(doc, footnote, 9, color=(0x40, 0x40, 0x40), after=0).paragraph_format.space_before = Pt(6)
    doc.save(path)
    print(f"wrote {path}")


def emit(name, title, subtitle, caption, headers, rows, aligns, tidy_df, bold_rows=(), footnote=None):
    os.makedirs(OUTDIR, exist_ok=True)
    tidy_df.to_csv(os.path.join(OUTDIR, f"{name}.csv"), index=False)
    print(f"wrote {os.path.join(OUTDIR, name)}.csv")
    render_png(os.path.join(OUTDIR, f"{name}.png"), title, subtitle, caption, headers, rows, aligns,
               bold_rows, footnote)
    build_docx(os.path.join(OUTDIR, f"{name}.docx"), title, subtitle, caption, headers, rows, aligns,
               bold_rows, footnote)


# ---------------------------------------------------------------- tables
def main():
    grid = pd.read_csv(find("glkn_grid_proportions_per_agent_5070_4agent.csv"))
    grid["area_km2"] = (grid["cell_side_m"] ** 2) / 1e6
    present_agents = [a for a in AGENTS if a in set(grid["agent"])]
    missing = [a for a in AGENTS if a not in present_agents]
    if missing:
        print(f"NOTE: grid csv missing agents: {missing}")

    sizes = (grid[["cell_side_px", "cell_side_m", "area_km2", "n_cells_complete"]]
             .drop_duplicates().sort_values("cell_side_px", ascending=False).reset_index(drop=True))
    # wide pivots of the per-agent counts, indexed by cell size
    wide = grid.pivot_table(index="cell_side_px", columns="agent", values="n_with_change")

    # --- table A: detection rate (%) by cell size, agents as columns ---
    a_headers = ["Cell (px)", "Side (m)", "Area (km²)"] + [DISPLAY[a] + "\n(%)" for a in present_agents]
    a_rows, a_tidy = [], []
    for _, s in sizes.iterrows():
        px = int(s.cell_side_px)
        base = [f"{px:,}", f"{int(s.cell_side_m):,}", f"{s.area_km2:,.2f}"]
        rates = []
        trow = {"cell_side_px": px, "cell_side_m": int(s.cell_side_m), "area_km2": round(s.area_km2, 4),
                "n_cells_complete": int(s.n_cells_complete)}
        for a in present_agents:
            nwc = wide.loc[px, a]
            rate = 100.0 * nwc / s.n_cells_complete
            rates.append(f"{rate:.1f}")
            trow[f"detection_rate_pct_{a}"] = round(rate, 1)
        a_rows.append(base + rates); a_tidy.append(trow)
    bold_idx = [i for i, s in sizes.iterrows() if int(s.cell_side_px) == SEL_PX]
    a_aligns = ["l", "r", "r"] + ["r" for _ in present_agents]
    emit("chapter3_table_detection_rate_by_cell_size",
         "Per-Agent Change Detection Rate by Grid Cell Size", None,
         "Fraction of complete cells, fully within the seven-watershed area of interest (AOI), that "
         "contain at least one polygon of each change agent, by grid cell size. Detection rate is "
         "n_with_change divided by n_cells_complete. The 112 pixel cell (11.29 square kilometers), the "
         "selected size, is shown in bold. GLKN change polygons, 2010 to 2020.",
         a_headers, a_rows, a_aligns, pd.DataFrame(a_tidy), bold_rows=bold_idx)

    # --- table B: absolute counts of complete cells with change, by cell size ---
    b_headers = ["Cell (px)", "Side (m)", "Area (km²)", "Complete\ncells"] + [DISPLAY[a] for a in present_agents]
    b_rows, b_tidy = [], []
    for _, s in sizes.iterrows():
        px = int(s.cell_side_px)
        base = [f"{px:,}", f"{int(s.cell_side_m):,}", f"{s.area_km2:,.2f}", f"{int(s.n_cells_complete):,}"]
        counts = [f"{int(wide.loc[px, a]):,}" for a in present_agents]
        trow = {"cell_side_px": px, "cell_side_m": int(s.cell_side_m), "area_km2": round(s.area_km2, 4),
                "n_cells_complete": int(s.n_cells_complete)}
        trow.update({f"n_with_change_{a}": int(wide.loc[px, a]) for a in present_agents})
        b_rows.append(base + counts); b_tidy.append(trow)
    b_aligns = ["l", "r", "r", "r"] + ["r" for _ in present_agents]
    emit("chapter3_table_complete_cell_counts_by_cell_size",
         "Complete Cells Containing Each Agent, by Grid Cell Size", None,
         "Absolute counts of complete cells intersecting each change agent's polygons, by grid cell "
         "size. The denominator for the detection rate is Complete cells (n_cells_complete), the number "
         "of cells fully within the seven-watershed AOI. The 112 pixel row is shown in bold. GLKN "
         "change polygons, 2010 to 2020.",
         b_headers, b_rows, b_aligns, pd.DataFrame(b_tidy), bold_rows=bold_idx)

    # --- table C: disturbance polygon size by agent, aggregated from the per-year change-agent files ---
    eda = pd.concat([pd.read_csv(f) for f in sorted(glob.glob(os.path.join(INDIR, "glkn_eda_changeagents_*.csv")))],
                    ignore_index=True)
    yrs = sorted(int(y) for y in eda["year"].unique())
    window = f"{yrs[0]} to {yrs[-1]}"
    # aggregate over the change years: n and total are sums, min and max are the extremes across years,
    # and mean is total area divided by n. the per-year files do not carry the underlying polygon areas,
    # so a pooled median and standard deviation are not recoverable and are not shown.
    agg = (eda.groupby("agent").agg(n_polys=("n_polys", "sum"), min_m2=("min_m2", "min"),
                                    max_m2=("max_m2", "max"), total_m2=("total_m2", "sum")).reset_index())
    agg["mean_m2"] = agg["total_m2"] / agg["n_polys"]
    agg = agg.sort_values("n_polys", ascending=False)
    c_headers = ["Agent", "N\npolygons", "Min\n(ha)", "Mean\n(ha)", "Max\n(ha)", "Total\n(km²)"]
    c_rows, c_tidy = [], []
    for _, r in agg.iterrows():
        ha = lambda col: r[col] / 1e4                                  # square meters to hectares
        c_rows.append([DISPLAY.get(r.agent, r.agent), f"{int(r.n_polys):,}",
                       f"{ha('min_m2'):,.2f}", f"{ha('mean_m2'):,.2f}", f"{ha('max_m2'):,.2f}",
                       f"{r.total_m2 / 1e6:,.2f}"])
        c_tidy.append({"agent": r.agent, "n_polys": int(r.n_polys), "min_ha": round(ha("min_m2"), 2),
                       "mean_ha": round(ha("mean_m2"), 2), "max_ha": round(ha("max_m2"), 2),
                       "total_km2": round(r.total_m2 / 1e6, 2)})
    c_aligns = ["l", "r", "r", "r", "r", "r"]
    foot = ("Polygon counts, extents, and totals aggregated over the change years " + window + ": N and "
            "total are summed, min and max are the extremes across years, and mean is total area divided "
            "by N. The per-year export does not carry the underlying polygon areas, so a pooled median "
            "and standard deviation are not recoverable and are not shown. Beaver and insect and disease "
            "rest on few polygons (see the N column), so their statistics are based on limited data.")
    emit("chapter3_table_polygon_size_by_agent",
         "Disturbance Polygon Size by Agent", f"GLKN watersheds, {window}",
         None, c_headers, c_rows, c_aligns, pd.DataFrame(c_tidy), footnote=foot)

    print("\nsummary:")
    print(f"  grid agents present: {present_agents}" + (f"  MISSING: {missing}" if missing else ""))
    print(f"  polygon-size agents (per-year files): {sorted(eda['agent'].unique())}, years {window}")


if __name__ == "__main__":
    main()
