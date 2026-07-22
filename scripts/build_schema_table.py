#!/usr/bin/env python3
"""Table 2.1, the ten-class classification schema, as an OSU-thesis-styled .docx plus a tidy CSV.

OSU Manuscript Format, not a journal: the caption sits above the table, the caption and label are the
same size as body text, and the table is numbered per chapter (2.1). Booktabs-style horizontal rules
(top, under the header, and at the foot), no vertical rules. Codes, class names, and stable/change
assignments are fixed; only the definition wording is lightly refined for parallelism.

Run: python scripts/build_schema_table.py
Requires: pandas, python-docx
"""

import os

import pandas as pd

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "manuscript_formatting", "schema_table")

CAPTION = ("Table 2.1. Ten-class classification schema. Change classes are defined by a transition "
           "observed between the two analysis dates rather than by an end-state land cover.")

# fixed codes, names, and types; definitions refined for parallelism, one concise line each, no
# em dashes, Oxford comma, ordered by code
ROWS = [
    (1, "Harvest", "Change", "Forest harvest or clearcut between the two dates"),
    (2, "Development", "Change", "Conversion to developed or built-up land between the two dates"),
    (3, "Stable forest", "Stable", "Forest, unchanged between the two dates"),
    (4, "Stable urban", "Stable", "Developed or built-up land, unchanged between the two dates"),
    (5, "Stable water", "Stable", "Open water, unchanged between the two dates"),
    (6, "Stable agriculture", "Stable", "Cultivated or agricultural land, unchanged between the two dates"),
    (7, "Stable grass/shrub", "Stable", "Grassland or shrubland, unchanged between the two dates"),
    (8, "Stable wetland", "Stable", "Wetland, unchanged between the two dates"),
    (9, "Beaver", "Change", "Beaver-induced flooding between the two dates"),
    (10, "Insect/disease mortality", "Change",
     "Forest mortality from insect or disease between the two dates"),
]
COLS = ["Code", "Class", "Type", "Definition"]
FONT = "Times New Roman"
SIZE = 12                                                  # body text size; caption and table match


def set_cell_border(cell, **edges):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
    for edge, spec in edges.items():
        el = tb.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}"); tb.append(el)
        for k, v in spec.items():
            el.set(qn(f"w:{k}"), str(v))


def build_docx(df, path):
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SIZE)

    # caption above the table, same size as body text
    cap = doc.add_paragraph()
    run = cap.add_run(CAPTION)
    run.font.name = FONT
    run.font.size = Pt(SIZE)

    thick = {"sz": 12, "val": "single", "color": "000000"}   # top and foot rules
    thin = {"sz": 6, "val": "single", "color": "000000"}     # under-header rule

    table = doc.add_table(rows=1, cols=len(COLS))
    table.autofit = True
    hdr = table.rows[0].cells
    for j, col in enumerate(COLS):
        p = hdr[j].paragraphs[0]
        r = p.add_run(col); r.bold = True; r.font.name = FONT; r.font.size = Pt(SIZE)
        set_cell_border(hdr[j], top=thick, bottom=thin)     # booktabs top and midrule
    for i, row in enumerate(df.itertuples(index=False)):
        cells = table.add_row().cells
        last = i == len(df) - 1
        for j, val in enumerate(row):
            p = cells[j].paragraphs[0]
            r = p.add_run(str(val)); r.font.name = FONT; r.font.size = Pt(SIZE)
            if last:
                set_cell_border(cells[j], bottom=thick)      # booktabs bottomrule
    doc.save(path)


def main():
    os.makedirs(OUT, exist_ok=True)
    df = pd.DataFrame(ROWS, columns=COLS)
    csv_path = os.path.join(OUT, "table_2_1_schema.csv")
    docx_path = os.path.join(OUT, "table_2_1_schema.docx")
    df.to_csv(csv_path, index=False)
    build_docx(df, docx_path)
    print(f"wrote {docx_path}")
    print(f"wrote {csv_path}")
    print(df.to_string(index=False))


if __name__ == "__main__":
    main()
