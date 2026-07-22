#!/usr/bin/env python3
"""Format existing repo results into manuscript-ready tables (tidy CSV + styled .docx).

This is a FORMATTING task. It reads numbers only from the results CSVs already in reports/ and
rounds them; it computes no new results and hand-types no values. Each table traces to its source
CSV, listed in the table note and in the printed manifest. Rendered for a Google Docs workflow:
tables go to one manuscript_tables.docx (booktabs-like styling, python-docx) and each table also
lands as a tidy CSV under manuscript_formatting/tables/ for traceability and paste-as-table.

Run: python scripts/build_manuscript_formatting.py
Requires: pandas, python-docx
"""

import os

import pandas as pd

R = "reports"
OUT = "manuscript_formatting"
TAB = os.path.join(OUT, "tables")

# canonical labels used across every table
SRC_LABEL = {"embedding_v2": "v2", "embedding_v3": "v3", "embedding_v4": "v4",
             "embedding_v5": "v5", "embedding_v6": "v6", "spectral_specall": "spec_all"}
SRC_ORDER = ["v2", "v3", "v4", "v5", "v6", "spec_all"]
CLS10_ORDER = ["harvest", "development", "forest", "urban", "water", "ag", "grass_shrub",
               "wetland", "beaver", "insect_disease"]
CLS10_DISPLAY = {"harvest": "Harvest", "development": "Development", "forest": "Forest",
                 "urban": "Urban", "water": "Water", "ag": "Agriculture",
                 "grass_shrub": "Grass/Shrub", "wetland": "Wetland", "beaver": "Beaver",
                 "insect_disease": "Insect/Disease"}
CLS5_ORDER = ["Stable", "Harvest", "Development", "Insect/Disease", "Beaver"]
ACC = 3          # decimals for accuracy-like quantities (OA, kappa, F1, IoU, UA, PA)


def r(x, n=ACC):
    return round(float(x), n) if pd.notna(x) else x


# ---- table builders: each returns dict(id, title, note, sources, df) -------------------------

def t_overall_10():
    d = pd.read_csv(f"{R}/spectral_composite_classified_maps/comparison/overall_comparison.csv")
    d = d[d.bracket == "pooled"].copy()
    d["source"] = d.source.map(SRC_LABEL)
    d = d.set_index("source").loc[SRC_ORDER].reset_index()
    out = pd.DataFrame({
        "Source": d.source,
        "OA": d.OA.map(r), "Kappa": d.kappa.map(r),
        "Macro-F1": d.macro_F1.map(r), "Mean IoU": d.mean_IoU.map(r),
        "N cells": d.n_cells.astype(int), "N pixels": d.total_pixels.astype(int)})
    note = ("Pooled across the five NAIP brackets, 10-class schema, adjudicated interpreted "
            "reference. Embedding variants v2 to v6 pool 180 cells; spec_all pools 168 cells, "
            "since 12 spec_all rasters are entirely nodata in the out-of-sample brackets, so the "
            "embedding and spectral pooled columns rest on different cell sets. Accuracies rounded "
            "to 3 decimals. Source: reports/spectral_composite_classified_maps/comparison/"
            "overall_comparison.csv (bracket = pooled).")
    return dict(id="table_2_3", title="Table 2.3. Pooled classification accuracy by prediction source "
                "(10-class schema).", note=note,
                sources=["overall_comparison.csv"], df=out)


def t_overall_10_bracket():
    d = pd.read_csv(f"{R}/spectral_composite_classified_maps/comparison/overall_comparison.csv")
    d["source"] = d.source.map(SRC_LABEL)
    piv = d.pivot(index="source", columns="bracket", values="OA").loc[SRC_ORDER]
    cols = ["2017_2019", "2018_2020", "2019_2021", "2020_2022", "2021_2023", "pooled"]
    piv = piv[cols].map(r).reset_index()
    piv.columns = ["Source", "2017-2019", "2018-2020 (control)", "2019-2021", "2020-2022",
                   "2021-2023", "Pooled"]
    # tidy long companion carries n_cells and kappa
    tidy = d[["source", "bracket", "control", "OA", "kappa", "n_cells", "total_pixels"]].copy()
    tidy["OA"] = tidy.OA.map(r); tidy["kappa"] = tidy.kappa.map(r)
    note = ("Overall accuracy (OA) per source and NAIP bracket, 10-class schema, adjudicated "
            "reference. The 2018-2020 bracket is the in-sample control (training window). Brackets "
            "use disjoint 36-cell sets, so a bracket-to-bracket difference confounds temporal "
            "transfer with cell composition, and these are five independent assessments, not a "
            "controlled transfer curve. Embeddings use 36 cells per bracket (180 pooled); spec_all "
            "uses 36, 36, 34, 32, and 30 cells across the brackets (168 pooled). OA rounded to 3 "
            "decimals. Source: overall_comparison.csv. The tidy CSV adds kappa and n_cells per "
            "bracket.")
    return dict(id="table_2_4", title="Table 2.4. Overall accuracy by source and temporal bracket "
                "(10-class schema).", note=note,
                sources=["overall_comparison.csv"], df=piv, tidy=tidy)


def t_perclass_f1_10():
    d = pd.read_csv(f"{R}/spectral_composite_classified_maps/comparison/combined_metrics_long.csv")
    d = d[(d.bracket == "pooled") & d.class_name.notna()].copy()
    d["source"] = d.source.map(SRC_LABEL)
    f1 = d.pivot(index="class_name", columns="source", values="f1").reindex(CLS10_ORDER)[SRC_ORDER]
    f1 = f1.map(r)
    # support differs between the 180-cell embeddings and the 168-cell spec_all
    sup_emb = d[d.source == "v2"].set_index("class_name").support.reindex(CLS10_ORDER).astype(int)
    sup_spec = d[d.source == "spec_all"].set_index("class_name").support.reindex(CLS10_ORDER).astype(int)
    out = f1.reset_index()
    out.insert(0, "Class", out.class_name.map(CLS10_DISPLAY)); out = out.drop(columns="class_name")
    out["Support (emb, 180)"] = sup_emb.values
    out["Support (spec, 168)"] = sup_spec.values
    note = ("Per-class F1 (pooled across brackets), 10-class schema, adjudicated reference. F1 = "
            "2*UA*PA/(UA+PA). Reference support is the reference pixel count for the class; it "
            "differs between the 180-cell embedding pool and the 168-cell spec_all pool, so both "
            "are given. Support is identical across v2 to v6 (shared reference). F1 rounded to 3 "
            "decimals. Full per-class UA, PA, F1, IoU, and support for every source are in the "
            "supplementary long table. Source: combined_metrics_long.csv (bracket = pooled).")
    return dict(id="table_2_5", title="Table 2.5. Per-class F1 by prediction source (10-class schema, "
                "pooled).", note=note, sources=["combined_metrics_long.csv"], df=out)


def t_overall_5():
    d = pd.read_csv(f"{R}/spectral_composite_classified_maps/collapsed_5class/comparison_collapsed.csv")
    d = d.set_index("source").loc[SRC_ORDER].reset_index()
    out = pd.DataFrame({
        "Source": d.source, "OA": d.OA.map(r),
        "All-Stable baseline OA": d.all_stable_baseline.map(r),
        "Kappa": d.kappa.map(r), "Macro-F1": d.macro_F1.map(r), "Mean IoU": d.mean_IoU.map(r),
        "N cells": d.n_cells.astype(int)})
    note = ("Pooled accuracy under the canonical 5-class collapse (Stable, Harvest, Development, "
            "Insect/Disease, Beaver), all six sources on the common 168-cell set. In this collapse "
            "the reference Other pixels fold into Stable and Unknown pixels are excluded, which is "
            "the single collapse now used across the whole 5-class analysis, including the other "
            "5-class exhibits (former Tables 6, 7, 10, and 11, now unassigned; see "
            "renumber_manifest.md). The all-Stable baseline OA is the accuracy of predicting Stable "
            "everywhere, "
            "near 0.985, so a high 5-class OA reflects the stable-dominated landscape and kappa and "
            "macro-F1 carry the change-class signal. Accuracies rounded to 3 decimals. Source: "
            "collapsed_5class/comparison_collapsed.csv.")
    return dict(id="T4", title="(unassigned) Pooled accuracy by prediction source (5-class collapse, "
                "common 168-cell set).", note=note,
                sources=["comparison_collapsed.csv"], df=out)


def _interp_ci(path, unit_col="n_pairs"):
    d = pd.read_csv(path)
    out = pd.DataFrame({
        "Class": d.cls, unit_col.replace("_", " ").title(): d[unit_col].astype(int),
        "Support (px)": d.support_px.astype(int),
        "F1": d.f1.map(r), "F1 95% CI": [f"{r(lo)}-{r(hi)}" for lo, hi in zip(d.f1_lo, d.f1_hi)],
        "IoU": d.iou.map(r), "IoU 95% CI": [f"{r(lo)}-{r(hi)}" for lo, hi in zip(d.iou_lo, d.iou_hi)],
        "Reliability": d.reliability})
    return out


def t_interp_10():
    out = _interp_ci(f"{R}/interpreter_agreement/per_class_agreement_ci.csv")
    note = ("Inter-interpreter per-class agreement F1 with 95% cluster (pair) bootstrap confidence "
            "intervals, 10-class schema, over the 72 double-interpreted cells (the resampling unit "
            "is the pair, not the pixel). F1 is the balanced probability that the two interpreters "
            "concur given one assigned the class. Reliability tiers on F1: High at or above 0.70, "
            "Moderate 0.50 to 0.70, Low below 0.50. Pairs is the number of pairs in which the class "
            "occurs. Values rounded to 3 decimals. Source: "
            "interpreter_agreement/per_class_agreement_ci.csv.")
    return dict(id="table_2_8", title="Table 2.8. Inter-interpreter per-class agreement with 95% CIs "
                "(10-class schema).", note=note,
                sources=["per_class_agreement_ci.csv"], df=out)


def t_interp_5():
    out = _interp_ci(f"{R}/interpreter_agreement/per_class_agreement_ci_5class.csv")
    note = ("Inter-interpreter per-class agreement F1 with 95% cluster (pair) bootstrap confidence "
            "intervals under the 5-class collapse, over the 72 double-interpreted cells. Same "
            "method as the 10-class agreement table. Interpreters agree almost perfectly on Stable "
            "and well on Harvest, and fall to Low reliability on Development, Insect/Disease, and "
            "Beaver, so a model scored on those classes against a single interpretation is bounded "
            "by reference noise. Values rounded to 3 decimals. Source: "
            "interpreter_agreement/per_class_agreement_ci_5class.csv.")
    return dict(id="T6", title="(unassigned) Inter-interpreter per-class agreement with 95% CIs "
                "(5-class collapse).", note=note,
                sources=["per_class_agreement_ci_5class.csv"], df=out)


def t_model_vs_interp_5():
    d = pd.read_csv(f"{R}/model_vs_interpreter_5class/model_per_class_ci_5class.csv")
    d["source"] = d.cls  # placeholder to avoid clobber
    m = pd.read_csv(f"{R}/model_vs_interpreter_5class/model_per_class_ci_5class.csv")
    f1 = m.pivot(index="cls", columns="source", values="f1").reindex(CLS5_ORDER)[SRC_ORDER].map(r)
    interp = pd.read_csv(f"{R}/interpreter_agreement/per_class_agreement_ci_5class.csv").set_index("cls")
    out = f1.reset_index().rename(columns={"cls": "Class"})
    out["Interpreter ceiling"] = [r(interp.loc[c].f1) if c in interp.index else None for c in CLS5_ORDER]
    # tidy long carries CIs and n_cells
    tidy = m[["source", "cls", "n_cells", "support_px", "f1", "f1_lo", "f1_hi",
              "iou", "iou_lo", "iou_hi", "reliability"]].copy()
    for c in ["f1", "f1_lo", "f1_hi", "iou", "iou_lo", "iou_hi"]:
        tidy[c] = tidy[c].map(r)
    note = ("Per-class F1 under the 5-class collapse for each prediction source against the "
            "adjudicated reference, with the inter-interpreter agreement F1 as the ceiling column. "
            "Model F1 is pooled over each source's usable cells (v2 to v6 on 180 cells, spec_all on "
            "168) with a cluster (cell) bootstrap; the interpreter ceiling uses the cluster (pair) "
            "bootstrap over 72 pairs. Point estimates rounded to 3 decimals; the tidy CSV carries "
            "the 95% CIs, IoU, and per-class cell counts. Sources: "
            "model_vs_interpreter_5class/model_per_class_ci_5class.csv and "
            "interpreter_agreement/per_class_agreement_ci_5class.csv.")
    return dict(id="T7", title="(unassigned) Per-class F1 by source versus the inter-interpreter "
                "ceiling (5-class collapse).", note=note,
                sources=["model_per_class_ci_5class.csv", "per_class_agreement_ci_5class.csv"],
                df=out, tidy=tidy)


def t_changecap():
    d = pd.read_csv(f"{R}/sensitivity_changecap/sensitivity_metrics_long.csv")
    d = d[d.is_change_class == True].copy()
    d["Class"] = d.class_name.map(CLS10_DISPLAY)
    d = d.sort_values(["class_name", "cap"])
    out = pd.DataFrame({
        "Change class": d.Class, "Training cap": d.cap.astype(int),
        "Training ceiling (px)": d.train_ceiling.astype(int),
        "UA (precision)": d.precision.map(r), "PA (recall)": d.recall.map(r), "F1": d.f1.map(r),
        "Support (px)": d.support.astype(int), "Predicted (px)": d.predicted_pixels.astype(int)})
    note = ("Training-cap sensitivity for the four change classes, 10-class schema, v2 embedding "
            "classifier trained on 2018/2020 with change classes capped at 50, 100, 150, and 200 "
            "training points (stable classes held at 200), pooled over the same 180 cells. UA is "
            "user's accuracy (precision), PA is producer's accuracy (recall). The training ceiling "
            "is the total unique training-pixel pool for the class, so the cap constrains beaver "
            "and insect_disease (small pools) far more than harvest and development. Support is the "
            "reference pixel count. Values rounded to 3 decimals. Source: "
            "sensitivity_changecap/sensitivity_metrics_long.csv.")
    return dict(id="table_2_7", title="Table 2.7. Training-cap sensitivity for the change classes "
                "(10-class schema, 180 cells).", note=note,
                sources=["sensitivity_metrics_long.csv"], df=out)


def t_spatial():
    d = pd.read_csv(f"{R}/spatial_structure/with_spec_all/spatial_structure_summary.csv")
    order = ["interpreted", "v2", "v3", "v4", "v5", "v6", "spec_all"]
    d = d.set_index("source").loc[[s for s in order if s in set(d.source)]].reset_index()
    out = pd.DataFrame({
        "Source": d.source.replace({"interpreted": "interpreted (ref)"}),
        "N patches": d.n_patches.astype(int),
        "Mean patch (ha)": d.mean_patch_ha.map(lambda x: round(x, 2)),
        "Median-by-area (ha)": d.median_by_area_ha.map(lambda x: round(x, 2)),
        "Moran's I (mean)": d.morans_i_mean.map(lambda x: round(x, 2)),
        "Moran's I (std)": d.morans_i_std.map(lambda x: round(x, 3))})
    note = ("Spatial-structure diagnostics per source, measured within the interpreted cell "
            "footprints so the reference sets the scale. Mean patch size and median-by-area (the "
            "patch size at which the cumulative area ECDF crosses 0.5) come from 8-connected "
            "component labeling; Moran's I is queen-contiguity spatial autocorrelation of the class "
            "raster, read as a smoothness diagnostic since class codes are nominal. The interpreted "
            "source is the adjudicated reviewer per cell; the model source is each cell's "
            "temporally-matched per-bracket prediction. Patch sizes rounded to 2 decimals, Moran's "
            "I to 2 (mean) and 3 (std) decimals. Source: "
            "spatial_structure/with_spec_all/spatial_structure_summary.csv.")
    return dict(id="table_2_6", title="Table 2.6. Spatial-structure diagnostics by source.", note=note,
                sources=["with_spec_all/spatial_structure_summary.csv"], df=out)


def t_percell_f1_5():
    d = pd.read_csv(f"{R}/per_cell_f1_5class/per_cell_f1_summary.csv")
    d = d.set_index("source").loc[SRC_ORDER].reset_index()
    out = pd.DataFrame({
        "Source": d.source,
        "N (common)": d.n_common.astype(int), "Mean F1 (common)": d.mean_common.map(r),
        "Median F1 (common)": d.median_common.map(r),
        "N (full)": d.n_full.astype(int), "Mean F1 (full)": d.mean_full.map(r),
        "Median F1 (full)": d.median_full.map(r)})
    note = ("Per-grid-cell macro-F1 under the 5-class collapse, each cell one sample, macro over "
            "the classes present in the reference or the prediction for that cell. The common set "
            "is the 168 cells usable for all six sources; each source's full set is its own usable "
            "cells (v2 to v6 have 180, spec_all 168), so the full-set columns are not a head-to-head "
            "ranking. Values rounded to 3 decimals. Source: "
            "per_cell_f1_5class/per_cell_f1_summary.csv.")
    return dict(id="T10", title="(unassigned) Per-cell 5-class macro-F1 summary by source.", note=note,
                sources=["per_cell_f1_summary.csv"], df=out)


def t_percell_change():
    d = pd.read_csv(f"{R}/per_cell_change_f1/change_f1_summary.csv")
    d["Change class"] = d.change_class.map({"harvest": "Harvest", "development": "Development",
                                            "insect_disease": "Insect/Disease", "beaver": "Beaver"})
    mean = d.pivot(index="change_class", columns="source", values="mean_f1")
    mean = mean.reindex(["harvest", "development", "insect_disease", "beaver"])[SRC_ORDER].map(r)
    mean = mean.reset_index()
    mean.insert(0, "Change class", mean.change_class.map({"harvest": "Harvest",
                "development": "Development", "insect_disease": "Insect/Disease", "beaver": "Beaver"}))
    mean = mean.drop(columns="change_class")
    tidy = d[["change_class", "source", "n_cells", "mean_f1", "median_f1", "frac_f1_zero"]].copy()
    for c in ["mean_f1", "median_f1", "frac_f1_zero"]:
        tidy[c] = tidy[c].map(r)
    note = ("Per-cell F1 for each change class and source under the 5-class collapse, each cell one "
            "sample, on the common 168-cell set. The table shows the mean per-cell F1; a cell "
            "contributes to a class only where that class is present in the reference or the "
            "prediction for that source, so the contributing cell count differs across sources and "
            "is given, with the median and the fraction of contributing cells with F1 = 0, in the "
            "tidy CSV. Values rounded to 3 decimals. Source: "
            "per_cell_change_f1/change_f1_summary.csv.")
    return dict(id="T11", title="(unassigned) Mean per-cell F1 by change class and source "
                "(5-class collapse, 168-cell common set).", note=note,
                sources=["change_f1_summary.csv"], df=mean, tidy=tidy)


# ---- supplementary tables --------------------------------------------------------------------

def s_perclass_full_10():
    d = pd.read_csv(f"{R}/spectral_composite_classified_maps/comparison/combined_metrics_long.csv")
    d = d[(d.bracket == "pooled") & d.class_name.notna()].copy()
    d["Source"] = d.source.map(SRC_LABEL)
    d["Class"] = d.class_name.map(CLS10_DISPLAY)
    d["Source"] = pd.Categorical(d.Source, SRC_ORDER, ordered=True)
    d["Class"] = pd.Categorical(d.Class, [CLS10_DISPLAY[c] for c in CLS10_ORDER], ordered=True)
    d = d.sort_values(["Source", "Class"])
    out = pd.DataFrame({
        "Source": d.Source, "Class": d.Class,
        "UA (precision)": d.precision.map(r), "PA (recall)": d.recall.map(r),
        "F1": d.f1.map(r), "IoU": d.iou.map(r),
        "Support (px)": d.support.astype(int), "Cells present": d.cells_present.astype(int)})
    note = ("Full per-class UA, PA, F1, IoU, and support for every source, pooled, 10-class schema, "
            "adjudicated reference. Embeddings pool 180 cells, spec_all 168. Cells present is the "
            "number of cells containing the class. Values rounded to 3 decimals. Source: "
            "combined_metrics_long.csv (bracket = pooled).")
    return dict(id="S1", title="Table S1. Full per-class metrics by source (10-class schema, "
                "pooled).", note=note, sources=["combined_metrics_long.csv"], df=out)


def s_dedup():
    d = pd.read_csv(f"{R}/model_comparison/dedup_sensitivity_summary.csv")
    note = ("Selection sensitivity of the earlier all-years model_comparison arm: the "
            "pick-one-interpretation-per-location draw repeated 100 times with different random "
            "selections, on the earlier 154-location snapshot. Reports the OA distribution per "
            "version. This robustness check is moot under the current adjudicated pipeline, where "
            "each cell already has exactly one chosen reviewer and there is nothing to resample, so "
            "it is retained only as an earlier-snapshot record and is not part of the current "
            "results. Values as stored. Source: model_comparison/dedup_sensitivity_summary.csv.")
    return dict(id="S2", title="Table S2. Dedup-selection sensitivity of overall accuracy "
                "(earlier 154-location snapshot).", note=note,
                sources=["dedup_sensitivity_summary.csv"], df=d)


def s_speckle():
    d = pd.read_csv(f"{R}/model_comparison_current/model_speckle.csv")
    out = pd.DataFrame({
        "Source": d.version,
        "Neighbor-change": d.neighbor_change.map(r),
        "Pooled OA": d.pooled_OA.map(r),
        "Valid pixel pairs": d.valid_pairs.astype("int64")})
    note = ("Map speckle per embedding variant on the current pipeline: neighbor-change is the "
            "fraction of horizontally-adjacent, both-valid pixel pairs whose class differs, computed "
            "over the current 180 per-bracket adjudicated cells (coverage 1.0, the cells are fully "
            "classified). Low is smooth, high is per-pixel speckle. Pooled OA is the current pooled "
            "overall accuracy against the adjudicated reference, so the table shows the speckle "
            "versus accuracy relationship. This replaces the earlier 154-location snapshot and "
            "matches Figure 2.9. Values rounded to 3 decimals. Source: "
            "model_comparison_current/model_speckle.csv.")
    return dict(id="S3", title="Table S3. Map speckle (neighbor-change) by embedding variant "
                "(current 180-cell pipeline).", note=note, sources=["model_speckle.csv"], df=out)


def s_design_5class_ci():
    d = pd.read_csv(f"{R}/collapsed_5class_confusion/metrics_long.csv")
    d = d[d.scope == "overall"].copy()
    piv = d.pivot(index="variant", columns="metric", values="estimate")
    keep = [c for c in ["OA", "kappa", "macro_F1", "mean_IoU", "baseline_OA"] if c in piv.columns]
    piv = piv[keep]
    idx = [v for v in ["v2", "v3", "v4", "v5", "v6", "spec_all"] if v in piv.index]
    piv = piv.loc[idx].map(r).reset_index().rename(columns={"variant": "Source"})
    note = ("Design-based pooled 5-class overall metrics with bootstrap uncertainty, from the "
            "adjudicated collapsed-confusion analysis. Reported here as point estimates; the source "
            "CSV carries the ratio and bootstrap CI bounds per metric and class. This is the "
            "CI-bearing 5-class companion to the flagged 5-class overall table (former Table 4); note "
            "its basis may differ (its own cell set), see the consistency report before combining "
            "them. "
            "Values rounded to 3 decimals. Source: collapsed_5class_confusion/metrics_long.csv.")
    return dict(id="S4", title="Table S4. Design-based pooled 5-class overall metrics (point "
                "estimates; CIs in source CSV).", note=note,
                sources=["collapsed_5class_confusion/metrics_long.csv"], df=piv)


# chapter 2 main-text tables, in new-number order (2.3 to 2.8)
CH2 = [t_overall_10, t_overall_10_bracket, t_perclass_f1_10, t_spatial, t_changecap, t_interp_10]
# kept but not assigned a chapter 2 number (the 5-class tables; the 5-class results appear as figures)
FLAGGED = [t_overall_5, t_interp_5, t_model_vs_interp_5, t_percell_f1_5, t_percell_change]
SUPP = [s_perclass_full_10, s_dedup, s_speckle, s_design_5class_ci]


# ---- docx rendering (booktabs-like) ----------------------------------------------------------

def _set_cell_border(cell, **edges):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
    for edge, spec in edges.items():
        tag = qn(f"w:{edge}")
        el = tb.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}"); tb.append(el)
        for k, v in spec.items():
            el.set(qn(f"w:{k}"), str(v))


def _add_table(doc, df):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ncol = len(df.columns)
    table = doc.add_table(rows=1, cols=ncol)
    table.autofit = True
    thick = {"sz": 12, "val": "single", "color": "000000"}
    thin = {"sz": 6, "val": "single", "color": "000000"}
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        p = hdr[j].paragraphs[0]
        run = p.add_run(str(col)); run.bold = True; run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        _set_cell_border(hdr[j], top=thick, bottom=thin)
    for i, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        last = i == len(df) - 1
        for j, col in enumerate(df.columns):
            v = row[col]
            txt = f"{v:,}" if isinstance(v, (int,)) or (hasattr(v, "dtype") and "int" in str(v.dtype)) else str(v)
            p = cells[j].paragraphs[0]
            run = p.add_run(txt); run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            if last:
                _set_cell_border(cells[j], bottom=thick)
    return table


def render_docx(ch2_specs, flagged_specs, supp_specs, captions_md_path, path):
    import docx
    from docx.shared import Pt, Inches
    doc = docx.Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)

    h = doc.add_heading("Chapter 2 tables and figure captions", level=0)
    doc.add_paragraph("Formatted from existing repository results, numbered in the thesis per-chapter "
                      "scheme (Table 2.n). Every table traces to a source CSV named in its note. "
                      "Numbers are rounded only, not recomputed.")

    def emit(specs, header):
        doc.add_heading(header, level=1)
        for spec in specs:
            doc.add_heading(spec["title"], level=2)
            _add_table(doc, spec["df"])
            n = doc.add_paragraph()
            run = n.add_run("Note. " + spec["note"]); run.italic = True; run.font.size = Pt(8)
            doc.add_paragraph("")

    emit(ch2_specs, "Chapter 2 main-text tables (Tables 2.1 to 2.8)")
    doc.add_paragraph("Table 2.1 (ten-class schema) is in tables/schema_table/; no Table 2.2 exists "
                      "yet. See renumber_manifest.md.")
    emit(flagged_specs, "Tables not assigned a Chapter 2 number (flagged; the 5-class results appear "
                        "in the main text as figures)")
    emit(supp_specs, "Supplementary tables (S-prefix retained)")

    # figure captions section, pasted from the markdown file
    doc.add_page_break()
    doc.add_heading("Draft figure captions", level=1)
    if os.path.exists(captions_md_path):
        for line in open(captions_md_path):
            line = line.rstrip("\n")
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                continue
            elif line.strip():
                doc.add_paragraph(line)
    doc.save(path)


def main():
    os.makedirs(TAB, exist_ok=True)
    specs_ch2 = [f() for f in CH2]
    specs_flagged = [f() for f in FLAGGED]
    specs_supp = [f() for f in SUPP]
    manifest = []
    for spec in specs_ch2 + specs_flagged + specs_supp:
        spec["df"].to_csv(os.path.join(TAB, f"{spec['id']}.csv"), index=False)
        if "tidy" in spec:
            spec["tidy"].to_csv(os.path.join(TAB, f"{spec['id']}_tidy.csv"), index=False)
        manifest.append((spec["id"], spec["title"].split(".", 1)[0], ", ".join(spec["sources"])))
    captions_md = os.path.join(OUT, "figure_captions.md")
    render_docx(specs_ch2, specs_flagged, specs_supp, captions_md,
                os.path.join(OUT, "manuscript_tables.docx"))
    print("wrote", os.path.join(OUT, "manuscript_tables.docx"))
    print(f"tidy CSVs -> {TAB}/ ({len(specs_ch2)} ch2 + {len(specs_flagged)} flagged + "
          f"{len(specs_supp)} supp)")
    for mid, title, src in manifest:
        print(f"  {mid:12} {title:46} <- {src}")


if __name__ == "__main__":
    main()
